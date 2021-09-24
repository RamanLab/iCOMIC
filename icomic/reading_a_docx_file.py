import docx

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.


def to_read(document1):
#    document = docx.Document('/data/anjana/bwa_params.docx')
    document = docx.Document(document1)
    table = document.tables[0]
    
    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    data = []
    needed = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
    
        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
    
        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
        
#    print(len(data))
    for i in range(21):
    # =============================================================================
    #     if data[i]['Essential'] == 'yes':
    #         print(data[i]['Short description'], data[i]['Default value'])
    # =============================================================================
        if "yes" in data[i]['Essential']:
#            stuff = data[i]['Short description'], data[i]['Default value']
#            print(data[i]['Short description'], data[i]['Default value'])
            needed.append(data[i]['Short description'])
#    print(len(needed))
    number = len(needed)
    
    return number

def using_number(x):
    for i in range(to_read(x)):
        print(i)
#    return y


path = '/data/Priyanka/other_pipelines/iCOMIC/BWA_MEM.docx'
print(using_number(path) )
    
            
        
# =============================================================================
#         count = data[i]['Essential'].count("yes")
#         print(count)
# =============================================================================
    
# =============================================================================
#         with open('testconfig', 'w') as tc:
#             tc.write(data[i]['Parameter'] + ': new value' )
#             tc.close()
# =============================================================================



